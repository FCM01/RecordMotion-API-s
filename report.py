from fileinput import filename
from math import prod
from msilib.schema import File
from pickletools import uint1
from platform import win32_edition
from docx  import Document
from docx.shared import Inches
import win32com.client
import datetime
import aspose.words as aw
from microfunction import tools
class report:

    def make_patient_percription(self,array,patient_name,patient_number,practise_name,doctor_name,phone_number,email):
        try:
            print("array=>",array)
            print(practise_name)
            print(patient_number)
            print(doctor_name)
            print(phone_number)
            print(email)
            today = datetime.date.today()
            document = Document()
            document.add_picture("./resources/FLOW2 Logo.png", width = Inches(1))
            document.add_heading("Perscription",0)
            p1 =document.add_paragraph("Dear ")
            p1.add_run(patient_name).bold=True
            p1.add_run()

            p2 =document.add_paragraph("Please find perscription for your recent doctors visit at :")
            p2.add_run(str(practise_name)).bold = True
            p2.add_run('.')
            [document.add_paragraph('')for _ in range(2)]

            table = document.add_table(rows=1 ,cols = 3)
            hdr_cells =table.rows[0].cells
            hdr_cells[0].text ="Medication Name"
            hdr_cells[1].text ="Quatity"
            hdr_cells[2].text ="Frequecy of use"

            for i in range(0,len(array)):
                hdr_cells[i].paragraphs[0].runs[0].font.bold =True
                rows_cells = table.add_row().cells
                rows_cells[0].text = array[i]["item_name"]
                rows_cells[1].text = array[i]["quantity"]
                rows_cells[2].text = array[i]["frequency_of_use"]


            [document.add_paragraph('')for _ in range(4)]

            document.add_paragraph("We appriciate your purchace")
            document.add_paragraph("Sincerely Yours:")
            document.add_paragraph(f"Dr {doctor_name}")
            document.add_heading("Contact Information", 1)
            document.add_heading("", 0)
            document.add_paragraph(f"Phone Number : {phone_number}  Email: {email}")
            document.save(f'./files/perscription-{patient_name}-{today}.docx')
            doc = report()
            filename  = f"doctors_note-{patient_name}-{today}"
            file_in=f"C:\\Users\\farai\\OneDrive\\Documents\\FLOW2\\API's\\files\\{filename}.docx"
            file_out=f"C:\\Users\\farai\\OneDrive\\Documents\\FLOW2\\API's\\files\\Pdfs\\{filename}.pdf"
            convert = doc.docx_to_pdf(file_in,file_out)
            if convert == True:
                print("Done")
            tool = tools()
            id_number  = tool.generate_task_id()
            resp = {"id":id_number,"type":"perscription","patient_number":f"{patient_number}","url":f'./files/perscription:{patient_name}-{today}.docx',"pdf_url":f"./files/Pdfs/perscription:{patient_name}-{today}.pdf","filename":f"perscription:{patient_name}-{today}.docx"}
            return resp
        except Exception as e:
            print("ERROR:",e)
    def make_doctors_note(self,patient_name,organisation_name,practise_name,patient_number,doctor_name,doctors_notes,phone_number,email):
        try:
            if organisation_name == "":
                organisation_name ="Employer/Company"
            today = datetime.date.today()
            document = Document()
            document.add_picture("./resources/FLOW2 Logo.png", width = Inches(1))
            document.add_heading("Doctors Note",0)
            p1 =document.add_paragraph("Dear ")
            p1.add_run(organisation_name).bold=True
            p1.add_run()
            p2 = document.add_paragraph("This is a doctors note for the patient ")
            p2.add_run(str(patient_name)).bold = True
            p2.add_run(' who recently visited the dotor at ')
            p2.add_run(str(practise_name)).bold = True
            [document.add_paragraph('') for _ in range(2)]
            document.add_heading("Note", 1)
            p3 = document.add_paragraph(f"{doctors_notes}")
            p3.add_run()

            [document.add_paragraph('')for _ in range(2)]

            document.add_paragraph("We appriciate your visit")
            document.add_paragraph("Sincerely Yours:")
            document.add_paragraph(f"Dr {doctor_name}")
            [document.add_paragraph('') for _ in range(2)]
            document.add_heading("Contact Information", 1)
            document.add_heading("", 0)
            document.add_paragraph(f"Phone Number : {phone_number}  Email: {email}")
            document.save(f'./files/doctors_note-{patient_name}-{today}.docx')
            doc = report()
            filename  = f"doctors_note-{patient_name}-{today}"
            file_in=f"C:\\Users\\farai\\OneDrive\\Documents\\FLOW2\\API's\\files\\{filename}.docx"
            file_out=f"C:\\Users\\farai\\OneDrive\\Documents\\FLOW2\\API's\\files\\Pdfs\\{filename}.pdf"
            convert = doc.docx_to_pdf(file_in,file_out)
            if convert == True:
                print("Done")
            tool = tools()
            id_number = tool.generate_task_id()
            resp = {"id": id_number,"type":"doctors_note" ,"patient_number": f"{patient_number}","url": f'./files/doctors_note:{patient_name}-{patient_number}-{today}.docx',"pdf_url":f"./files/Pdfs/doctors_note:{patient_name}-{patient_number}-{today}.pdf" ,"filename": f"doctors_note-{patient_name}-{today}"}
            return resp
        except Exception as e:
            print("ERROR:",e)
    def docx_to_pdf(self,src,dist):
        try:
            word = win32com.client.Dispatch("Word.Application")
            wdFormatPDF = 17
            doc  = word.Documents.Open(src)
            print(doc)
            doc.SaveAs(dist,FileFormat = wdFormatPDF)
            doc.Close()
            word.Quit()
            return True
        except Exception as e :
            print("ERROR",e)

#


